"""Generate RIS MVP BPMN + Forms + Integration DOCX."""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def inline_md(para, text):
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'; r.font.size = Pt(9.5)
        else:
            para.add_run(part)

# ── document setup ─────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(1.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

BLUE  = RGBColor(0x1F, 0x49, 0x7D)
SIZES = {1: 16, 2: 14, 3: 12, 4: 11}
for lvl in range(1, 5):
    h = doc.styles[f'Heading {lvl}']
    h.font.name = 'Calibri'
    h.font.color.rgb = BLUE
    h.font.size = Pt(SIZES[lvl])

ACCENT  = '2E74B5'
STRIPE  = 'D6E4F0'
WHITE   = 'FFFFFF'
GREEN   = '1F7A4D'   # for integration header rows
LBLUE   = 'EAF3FB'

# ── cover ──────────────────────────────────────────────────────────────────────

doc.add_paragraph().add_run('\n\n')

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('RIS MVP')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run('Бизнес-процессы · Формы · Интеграционный задел')
r2.bold = True; r2.font.size = Pt(15)

doc.add_paragraph()
mp = doc.add_paragraph()
mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mp.add_run('Версия 1.0  |  2026-06-13  |  Республика Казахстан').font.size = Pt(12)
doc.add_page_break()

# ── parse ──────────────────────────────────────────────────────────────────────

def is_sep(line):
    return bool(re.match(r'^\s*\|?\s*[-:]+[-| :]*\s*$', line))

def parse_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]

def flush_table(rows):
    if not rows: return
    cols = max(len(r) for r in rows)
    padded = [r + [''] * (cols - len(r)) for r in rows]
    tbl = doc.add_table(rows=len(padded), cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    for ri, row in enumerate(padded):
        for ci, txt in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            if ri == 0:
                set_cell_bg(cell, ACCENT)
                r = para.add_run(txt)
                r.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
            else:
                set_cell_bg(cell, STRIPE if ri % 2 == 0 else WHITE)
                inline_md(para, txt)
                for run in para.runs: run.font.size = Pt(10)
    doc.add_paragraph()

def flush_code(lines):
    text = ''.join(lines).rstrip('\n')
    if not text: return
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(3)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF3F8')
    pPr.append(shd)
    r = para.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(8.5)
    doc.add_paragraph()

# ── main parse ─────────────────────────────────────────────────────────────────

with open('ris-mvp-bpmn.md', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False; in_code = False
table_rows = []; code_lines = []

i = 0
while i < len(lines):
    line = lines[i]; s = line.rstrip('\n')

    if s.startswith('```'):
        if in_table: flush_table(table_rows); table_rows = []; in_table = False
        if in_code:  flush_code(code_lines);  code_lines = []; in_code = False
        else:        in_code = True
        i += 1; continue

    if in_code:
        code_lines.append(line); i += 1; continue

    if s.startswith('|') or (in_table and s.startswith('|')):
        if is_sep(s): i += 1; continue
        table_rows.append(parse_row(s))
        in_table = True; i += 1; continue

    if in_table:
        flush_table(table_rows); table_rows = []; in_table = False

    if not s or s == '---':
        doc.add_paragraph(); i += 1; continue

    m = re.match(r'^(#{1,4})\s+(.*)', s)
    if m:
        doc.add_heading(m.group(2), level=min(len(m.group(1)), 4))
        i += 1; continue

    mb = re.match(r'^(\s*)[-*]\s+(.*)', s)
    mn = re.match(r'^(\s*)\d+\.\s+(.*)', s)
    if mb or mn:
        match = mb or mn
        para = doc.add_paragraph(style='List Bullet' if mb else 'List Number')
        if len(match.group(1)) >= 4:
            para.paragraph_format.left_indent = Cm(1.5)
        inline_md(para, match.group(2))
        for run in para.runs: run.font.size = Pt(11)
        i += 1; continue

    para = doc.add_paragraph()
    inline_md(para, s)
    para.paragraph_format.space_after = Pt(4)
    i += 1

if in_table: flush_table(table_rows)
if in_code:  flush_code(code_lines)

doc.save('ris-mvp-bpmn.docx')
print('Saved: ris-mvp-bpmn.docx')
