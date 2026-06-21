"""Generate RIS MVP TZ as DOCX from the markdown source."""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top/bottom/left/right = color hex or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in kwargs.items():
        border = OxmlElement(f'w:{side}')
        if color:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
        else:
            border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_run_bold(para, text):
    run = para.add_run(text)
    run.bold = True
    return run

def add_run_code(para, text, doc):
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return run


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(1.5)

# Default styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)
doc.styles['Heading 4'].font.size = Pt(11)


# ── cover page ────────────────────────────────────────────────────────────────

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run('\n\n\n')

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run('ТЕХНИЧЕСКОЕ ЗАДАНИЕ')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_para.add_run('MVP радиологической информационной системы (RIS)')
r2.bold = True; r2.font.size = Pt(16)

doc.add_paragraph()
meta_para = doc.add_paragraph()
meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_para.add_run('Версия 2.0  |  Дата: 2026-06-13\nРеспублика Казахстан\n\n').font.size = Pt(12)

doc.add_page_break()


# ── read markdown ─────────────────────────────────────────────────────────────

with open('ris-mvp-tz-v2.md', encoding='utf-8') as f:
    lines = f.readlines()

# ── parser state ──────────────────────────────────────────────────────────────

in_table  = False
in_code   = False
table_rows = []
code_lines = []

ACCENT  = '2E74B5'   # header cell bg
STRIPE  = 'D6E4F0'   # alt row bg
WHITE   = 'FFFFFF'

SKIP_PATTERNS = [
    r'^\s*\|?\s*---+',   # separator rows
    r'^---\s*$',         # horizontal rule
    r'^\s*$',            # blank inside table parse
]

def is_separator(line):
    return bool(re.match(r'^\s*\|?\s*[-:]+[-| :]*\s*$', line))

def parse_row(line):
    line = line.strip().strip('|')
    return [c.strip() for c in line.split('|')]

def inline_md(para, text):
    """Render **bold**, `code`, and plain text into a paragraph."""
    # split on ** and `
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
        else:
            para.add_run(part)

def flush_table():
    global table_rows, in_table
    if not table_rows:
        in_table = False
        return

    col_count = max(len(r) for r in table_rows)
    # pad rows
    padded = [r + [''] * (col_count - len(r)) for r in table_rows]

    tbl = doc.add_table(rows=len(padded), cols=col_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    for ri, row in enumerate(padded):
        tr = tbl.rows[ri]
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

            if ri == 0:
                set_cell_bg(cell, ACCENT)
                r = para.add_run(cell_text)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
            else:
                bg = STRIPE if ri % 2 == 0 else WHITE
                set_cell_bg(cell, bg)
                para.paragraph_format.left_indent = Pt(2)
                inline_md(para, cell_text)
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    table_rows = []
    in_table   = False

def flush_code():
    global code_lines, in_code
    text = ''.join(code_lines).rstrip('\n')
    if text:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(4)
        # grey shading
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EEF3F8')
        pPr.append(shd)
        r = para.add_run(text)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
    code_lines = []
    in_code    = False


# ── main parse loop ───────────────────────────────────────────────────────────

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.rstrip('\n')

    # ── code block ──
    if stripped.startswith('```'):
        if in_table:
            flush_table()
        if in_code:
            flush_code()
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # ── table row ──
    if stripped.startswith('|') or (in_table and stripped.startswith('|')):
        if is_separator(stripped):
            i += 1
            continue
        row = parse_row(stripped)
        if not in_table:
            in_table = True
        table_rows.append(row)
        i += 1
        continue

    if in_table:
        flush_table()

    # ── skip blank / horizontal rule ──
    if not stripped or stripped == '---':
        if not in_table:
            doc.add_paragraph()
        i += 1
        continue

    # ── headings ──
    m = re.match(r'^(#{1,4})\s+(.*)', stripped)
    if m:
        level = len(m.group(1))
        text  = m.group(2)
        para  = doc.add_heading(text, level=min(level, 4))
        i += 1
        continue

    # ── bullet / numbered list ──
    m_bullet = re.match(r'^(\s*)[-*]\s+(.*)', stripped)
    m_num    = re.match(r'^(\s*)\d+\.\s+(.*)', stripped)
    if m_bullet or m_num:
        match = m_bullet or m_num
        indent = len(match.group(1))
        text   = match.group(2)
        style  = 'List Bullet' if m_bullet else 'List Number'
        para   = doc.add_paragraph(style=style)
        if indent >= 4:
            para.paragraph_format.left_indent = Cm(1.5)
        inline_md(para, text)
        for run in para.runs:
            run.font.size = Pt(11)
        i += 1
        continue

    # ── normal paragraph ──
    para = doc.add_paragraph()
    inline_md(para, stripped)
    para.paragraph_format.space_after = Pt(4)
    i += 1

if in_table:
    flush_table()
if in_code:
    flush_code()


# ── save ──────────────────────────────────────────────────────────────────────

out = 'ris-mvp-tz-v2.docx'
doc.save(out)
print(f'Saved: {out}')
